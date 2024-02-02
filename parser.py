import docx
import lxml.etree as ET


class Node:
    def __init__(self, id, tag, name, parent=None, value=None):
        self.id = id
        self.tag = tag
        self.name = name
        self.parent = parent
        self.value = value
        self.children = []


def parse_content_description_to_tree(doc):

    tocs = (styles['toc 1'], styles['toc 2'], styles['toc 3'],)


    root = Node(0, 'main', 'ROOT')
    first_level_id = 1

    in_content_description = False
    for i in range(len(doc.paragraphs)):
        para = doc.paragraphs[i]

        if in_content_description:
            s_name = para.style.name

            if s_name not in tocs:
                return root

            if s_name == styles['toc 1']:
                node = Node(first_level_id, 'sec', ' '.join(para.text.split()[:-1]), root)
                root.children.append(node)
                first_level_id += 1
            else:
                id = para.text.split()[0].split('.')
                parent = root
                for i in range(len(id) - 1):
                    for ch in parent.children:
                        if ch.id == id[i]:
                            parent = ch
                            break
                node = Node(id[-1], 'subsec', ' '.join(para.text.split()[:-1]), parent)
                parent.children.append(node)

        if para.style.name == styles['content_description']:
            in_content_description = True


def parse_content_description_to_etree(doc):

    tocs = (styles['toc 1'], styles['toc 2'], styles['toc 3'],)


    root = ET.Element('content_description')
    last_first_level = None
    last_second_level = None

    in_content_description = False
    for i in range(len(doc.paragraphs)):
        para = doc.paragraphs[i]

        if in_content_description:
            s_name = para.style.name

            if s_name not in tocs:
                return root

            if s_name == styles['toc 1']:
                node = ET.SubElement(root, 'sec')
                name = ' '.join(para.text.split()[:-1])
                if name[0].isdigit():
                    name = ' '.join(name.split()[1:])
                node.set('name', name)

                last_first_level = node

            if s_name == styles['toc 2']:
                node = ET.SubElement(last_first_level, 'subsec1')
                name = ' '.join(para.text.split()[:-1])
                if name[0].isdigit():
                    name = ' '.join(name.split()[1:])
                node.set('name', name)

                last_second_level = node

            if s_name == styles['toc 3']:
                node = ET.SubElement(last_second_level, 'subsec2')
                name = ' '.join(para.text.split()[:-1])
                if name[0].isdigit():
                    name = ' '.join(name.split()[1:])
                node.set('name', name)

        if para.style.name == styles['content_description']:
            in_content_description = True
    return root


def parse_em_all(doc):
    root = parse_content_description_to_etree(doc)
    prev = root

    tocs = (styles['toc 1'], styles['toc 2'], styles['toc 3'],)

    levels = {
        'sec': 1,
        'subsec1': 2,
        'subsec2': 3,
    }

    doc_index = 0

    for elem in root.iter():
        if elem == root:
            continue

        if prev == root or levels[prev.tag] < levels[elem.tag]:
            prev = elem
            continue

        print(elem.get('name'))
        doc_index = 0

        while doc_index < len(doc.paragraphs):
            if prev.get('name') in doc.paragraphs[doc_index].text and doc.paragraphs[doc_index].style.name not in tocs:
                in_iter = False
                s_name = doc.paragraphs[doc_index].style.name
                iter_start = prev
                iter1_start = prev
                numpstart = prev
                doc_index += 1
                for i in range(doc_index, len(doc.paragraphs)):
                    para = doc.paragraphs[i]
                    print(para.style.name)
                    if para.style.name == s_name:
                        doc_index = i
                        break

                    if para.style.name == styles['nump 2']:
                        node = ET.SubElement(prev, 'nump2')
                        node.text = para.text
                        numpstart = node
                        continue
                    elif para.style.name == styles['nump 3']:
                        node = ET.SubElement(prev, 'nump3')
                        node.text = para.text
                        numpstart = node
                        continue
                    else:
                        numpstart = prev

                    if para.style.name == styles['iter 1']:
                        if not in_iter:
                            in_iter = True
                            iter_start = ET.SubElement(prev, 'iter')
                        node = ET.SubElement(iter_start, 'i1')
                        node.text = para.text
                        iter1_start = node
                        continue

                    if para.style.name == styles['iter 2']:
                        node = ET.SubElement(iter1_start, 'i2')
                        node.text = para.text
                        continue

                    if para.style.name == styles['iter alt']:
                        if not in_iter:
                            in_iter = True
                            iter_start = ET.SubElement(prev, 'iter alt')
                        if iter_start.tag == 'iter alt':
                            node = ET.SubElement(iter_start, 'i1')
                        else:
                            node = ET.SubElement(iter1_start, 'i2')
                        node.text = para.text
                        continue

                    in_iter = False

                    if para.style.name == styles['normal']:
                        node = ET.SubElement(numpstart, 'n')
                        node.text = para.text
                        continue

                    if styles['extra'] in para.style.name:
                        node = ET.SubElement(prev, 'extra')
                        node.text = para.text
                        numpstart = node

                    node = ET.SubElement(prev, para.style.name[:3])
                    node.text = para.text
                break
            doc_index += 1
        prev = elem
    # doc_index = 0
    # while doc_index < len(doc.paragraphs):
    #     if prev.get('name') in doc.paragraphs[doc_index].text and doc.paragraphs[doc_index].style.name not in tocs:
    #         s_name = doc.paragraphs[doc_index].style.name
    #         doc_index += 1
    #         for i in range(doc_index, len(doc.paragraphs)):
    #             para = doc.paragraphs[i]
    #             print(para.style.name)
    #
    #             if para.style.name == s_name:
    #                 doc_index = i
    #                 break
    #             if para.style.name == styles['iter 1']:
    #                 node = ET.SubElement(prev, 'i1')
    #                 node.text = para.text
    #                 iter_start = node
    #
    #             elif para.style.name == styles['iter 1 alt']:
    #                 node = ET.SubElement(prev, 'i1alt')
    #                 node.text = para.text
    #                 iter_start = node
    #
    #             elif para.style.name == styles['iter 2']:
    #                 node = ET.SubElement(iter_start, 'i2')
    #                 node.text = para.text
    #
    #             else:
    #                 node = ET.SubElement(prev, 'p')
    #                 node.text = para.text
    #         break
    #     doc_index += 1
    return root

styles = {
    'content_description': 'Название-caps',  # Содержание
    'toc 1': 'toc 1',
    'toc 2': 'toc 2',
    'toc 3': 'toc 3',  # Перечисления в содержании (индексация длины 1, 2 и 3 соответственно)
    'subsec v': 'Нумерованный заголовок 3',
    'iter 1': "Перечисление-1",
    'iter 2': "Перечисление 2 уровень",
    'iter alt': "Перечень",
    'nump 2': "Нумерованный абзац 2",
    'nump 3': "Нумерованный абзац 3",
    'normal': "Normal",
    'extra': "Приложение",
}


def run_parser(doc):
    root = parse_em_all(doc)

    et = ET.ElementTree(root)
    et.write('output.xml', encoding="utf-8", pretty_print=True)

    root = parse_content_description_to_etree(doc)
    et = ET.ElementTree(root)
    et.write('output2.xml', encoding="utf-8", pretty_print=True)


doc = docx.Document("input.docx")
run_parser(doc)
